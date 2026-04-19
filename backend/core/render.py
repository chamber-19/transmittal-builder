"""
Core rendering logic for transmittal .docx generation.

Ported from Suite's backend/Transmittal-Builder/core/transmittal_render.py.
Standalone — no GUI dependencies, no Suite imports.
"""

from __future__ import annotations

import os
import re
from typing import Dict, List, Optional

from docx import Document
from docx.table import Table

# ─── Filename Parsing ─────────────────────────────────────────

_H = r"[-–—]"
DOC_ID_RE = re.compile(rf"(R3P{_H}(\d+){_H}E(\d+){_H}(\d+))", re.IGNORECASE)

# Matches a single leading "XMTL" prefix segment (with optional dash/
# underscore/space and Unicode dash variants). Applied iteratively in
# _normalize_xmtl_num so pathological inputs like "XMTL-XMTL-001" are
# fully stripped without using a nested quantifier (which would be a
# polynomial-ReDoS risk on whitespace-heavy inputs).
_XMTL_PREFIX_RE = re.compile(r"^xmtl[-_\s\u2013\u2014]*", re.IGNORECASE)


def _normalize_xmtl_num(raw: str) -> str:
    """
    Strip any user-supplied "XMTL"/"xmtl-"/"XMTL_"/"XMTL " prefix from a
    transmittal number. Returns the bare number (e.g. "001"). Empty input
    returns empty string. Mirror of the frontend `stripXmtlPrefix` helper.

    Iterative (rather than `(?:…)+`) to avoid catastrophic backtracking on
    inputs with long whitespace runs.
    """
    if not raw:
        return ""
    s = str(raw).strip()
    while True:
        new_s = _XMTL_PREFIX_RE.sub("", s).strip()
        if new_s == s:
            return s
        s = new_s


def format_xmtl_label(raw: str) -> str:
    """Return canonical "XMTL-<num>" label, never doubling the prefix."""
    bare = _normalize_xmtl_num(raw)
    return f"XMTL-{bare}" if bare else ""


def extract_doc_meta(filename: str) -> dict:
    """Parse a document filename into doc_no, desc, rev."""
    base = os.path.splitext(os.path.basename(filename))[0]
    m = DOC_ID_RE.search(base)
    if not m:
        return {"doc_no": "", "desc": base.strip(), "rev": ""}
    raw_doc = m.group(1)
    doc_no = re.sub(r"[–—]", "-", raw_doc).upper()
    pruned_doc_no = re.sub(r"^R3P-\d+-", "", doc_no)
    remainder = base[m.end(1):]
    remainder = re.sub(r"^[\s\-_–—:;|]+", "", remainder)
    return {"doc_no": pruned_doc_no, "desc": remainder.strip(), "rev": ""}


def _norm_key(doc_no: str) -> Optional[str]:
    """Normalize a doc number to E{n}-{nnnn} for revision map lookups."""
    m = re.search(r"E(\d+)[\-_–—](\d+)", doc_no or "", re.IGNORECASE)
    if not m:
        return None
    return f"E{int(m.group(1))}-{int(m.group(2)):04d}"


# ─── Word Document Helpers ────────────────────────────────────

def _walk_paragraphs(doc: Document):
    """Yield every paragraph in the doc, including nested tables and headers/footers."""
    def _iter(container):
        if hasattr(container, "paragraphs"):
            for p in container.paragraphs:
                yield p
        if hasattr(container, "tables"):
            for t in container.tables:
                for row in t.rows:
                    for cell in row.cells:
                        yield from _iter(cell)

    yield from _iter(doc)
    for section in doc.sections:
        yield from _iter(section.header)
        yield from _iter(section.footer)


def _replace_in_runs(runs, needle: str, replacement: str) -> bool:
    """Replace text that may span multiple runs, preserving formatting."""
    if not needle:
        return False
    full = "".join(r.text for r in runs)
    idx = full.find(needle)
    if idx < 0:
        return False

    before, after = full[:idx], full[idx + len(needle):]
    for r in runs:
        r.text = ""
    chunks = [before, replacement, after]
    ri = 0
    for chunk in chunks:
        if not chunk:
            continue
        if ri >= len(runs):
            runs[-1].text += chunk
        else:
            runs[ri].text = chunk
            ri += 1
    return True


def replace_text_everywhere(doc: Document, mapping: dict) -> Dict[str, str]:
    """Replace placeholder text throughout the document."""
    replacements = {}
    for p in _walk_paragraphs(doc):
        if not p.runs:
            continue
        full_text = "".join(r.text for r in p.runs)
        for key, val in mapping.items():
            if key in full_text:
                if _replace_in_runs(p.runs, key, val):
                    replacements[key] = val
    return replacements


def set_checkbox(doc: Document, label: str, checked: bool) -> bool:
    """Toggle a checkbox glyph (☐/☑/☒) next to the given label."""
    label_lc = (label or "").strip().lower()
    if not label_lc:
        return False

    def _find_glyphs(cell):
        runs = []
        for p in cell.paragraphs:
            for r in p.runs:
                if any(sym in r.text for sym in ("☐", "☑", "☒")):
                    runs.append(r)
        for t in cell._element.xpath(".//w:t"):
            txt = t.text or ""
            if any(sym in txt for sym in ("☐", "☑", "☒")):
                runs.append(t)
        return runs

    def _scan(tbl: Table) -> bool:
        for row in tbl.rows:
            label_col = None
            box_cols = []
            for ci, cell in enumerate(row.cells):
                text = " ".join(p.text for p in cell.paragraphs).strip().lower()
                if label_lc in text:
                    label_col = ci
                if any(sym in text for sym in ("☐", "☑", "☒")):
                    box_cols.append(ci)

            if label_col is not None:
                chosen = None
                if any(b < label_col for b in box_cols):
                    chosen = max(b for b in box_cols if b < label_col)
                elif box_cols:
                    chosen = min(box_cols)
                else:
                    for tc in [label_col - 1, label_col]:
                        if 0 <= tc < len(row.cells) and _find_glyphs(row.cells[tc]):
                            chosen = tc
                            break

                if chosen is not None:
                    for r in _find_glyphs(row.cells[chosen]):
                        if hasattr(r, "text"):
                            if checked:
                                r.text = r.text.replace("☐", "☒").replace("☑", "☒")
                            else:
                                r.text = r.text.replace("☑", "☐").replace("☒", "☐")
                            return True

            for cell in row.cells:
                for nested in cell.tables:
                    if _scan(nested):
                        return True
        return False

    for tbl in doc.tables:
        if _scan(tbl):
            return True
    for section in doc.sections:
        if section.header:
            for tbl in section.header.tables:
                if _scan(tbl):
                    return True
    return False


def find_table_by_headers(doc: Document, headers: List[str]) -> Optional[Table]:
    """Find a table whose first row contains all the given header strings."""
    for table in doc.tables:
        if not table.rows:
            continue
        row_text = [cell.text.strip() for cell in table.rows[0].cells]
        if all(h in row_text for h in headers):
            return table
    return None


def clear_table_body(table: Table) -> None:
    """Remove all rows except the header row."""
    while len(table.rows) > 1:
        table._element.remove(table.rows[-1]._element)


def append_rows(table: Table, items: List[dict]) -> None:
    """Append doc_no / desc / rev rows to a table."""
    for item in items:
        row = table.add_row()
        row.cells[0].text = item.get("doc_no", "")
        row.cells[1].text = item.get("desc", "")
        row.cells[2].text = item.get("rev", "")


def fill_contacts_table(doc: Document, contacts: List[Dict[str, str]]) -> None:
    """Fill the Name/Company/Email/Phone contacts table."""
    table = find_table_by_headers(doc, ["Name", "Company", "Email", "Phone"])
    if not table:
        return
    clear_table_body(table)
    for c in contacts:
        if any(c.values()):
            row = table.add_row()
            row.cells[0].text = c.get("name", "")
            row.cells[1].text = c.get("company", "")
            row.cells[2].text = c.get("email", "")
            row.cells[3].text = c.get("phone", "")


# ─── Core Render ──────────────────────────────────────────────

# All checkbox keys the template supports
CHECKBOX_MAP = [
    ("PDF", "trans_pdf"), ("CAD", "trans_cad"), ("Originals", "trans_originals"),
    ("Email", "via_email"), ("FTP", "via_ftp"),
    ("For Information Only", "ci_info"), ("For Approval", "ci_approval"),
    ("For Bid", "ci_bid"), ("For Preliminary", "ci_preliminary"),
    ("For Construction", "ci_const"), ("For As-Built", "ci_asbuilt"),
    ("For Fabrication", "ci_fab"), ("For Record", "ci_record"),
    ("For Reference", "ci_ref"),
    ("Approved", "vr_approved"), ("Approved as Noted", "vr_approved_noted"),
    ("Rejected", "vr_rejected"),
]


def _remove_reference_section(doc: Document) -> bool:
    """
    Remove the "REFERENCE" heading paragraph and its immediately-following
    table from the document. Used when no source PDFs are attached, so the
    rendered transmittal doesn't show an empty Reference table.

    Returns True if a section was removed.

    Detection: a paragraph whose text (stripped, lower-cased) starts with
    "reference" (covers "REFERENCE:", "Reference", "References", etc.).
    The first table that appears after that paragraph in document order is
    treated as the reference table and removed alongside the heading.
    """
    body = doc.element.body
    # qn() is needed to reference WordprocessingML element tags
    from docx.oxml.ns import qn
    p_tag = qn("w:p")
    tbl_tag = qn("w:tbl")

    # Helper to read concatenated text from a <w:p> element
    def _p_text(el) -> str:
        return "".join(t.text or "" for t in el.iter(qn("w:t")))

    target_p = None
    for child in list(body):
        if child.tag == p_tag:
            txt = _p_text(child).strip().lower().rstrip(":").strip()
            if txt.startswith("reference"):
                target_p = child
                break

    if target_p is None:
        return False

    # Walk forward from target_p, collecting siblings to remove until (and
    # including) the first <w:tbl>. If no table is found before the next
    # heading-like paragraph, just remove the heading itself.
    to_remove = [target_p]
    sibling = target_p.getnext()
    table_found = False
    while sibling is not None:
        to_remove.append(sibling)
        if sibling.tag == tbl_tag:
            table_found = True
            break
        # Stop if we hit another heading-style paragraph (defensive — keeps
        # us from eating subsequent sections if the template has no table)
        if sibling.tag == p_tag:
            txt = _p_text(sibling).strip()
            if txt and len(txt) < 40 and txt.isupper():
                # Looks like another heading — back off, only remove the
                # heading paragraph itself.
                to_remove = [target_p]
                break
        sibling = sibling.getnext()

    if not table_found and len(to_remove) > 1:
        # No table located — only safe to drop the heading itself
        to_remove = [target_p]

    for el in to_remove:
        body.remove(el)
    return True


def render_transmittal(
    template_path: str,
    fields: dict,
    checks: dict,
    contacts: List[Dict[str, str]],
    documents: List[Dict[str, str]],
    out_path: str,
    has_attached_drawings: bool = True,
) -> str:
    """
    Render a transmittal .docx from a template.

    Args:
        template_path: Path to the .docx template
        fields: Project/sender field values
        checks: Checkbox key → bool mapping
        contacts: List of {name, company, email, phone}
        documents: List of {doc_no, desc, rev}
        out_path: Where to save the rendered .docx
        has_attached_drawings: When False, the "REFERENCE" heading and its
            table are stripped from the output so no empty placeholder
            section is left behind.

    Returns:
        The output path.
    """
    if not os.path.isfile(template_path):
        raise FileNotFoundError("Template not found.")

    doc = Document(template_path)

    # Normalize the transmittal number so the rendered text always reads
    # "XMTL-001" — never "XMTL-XMTL-001" if the user already typed the prefix.
    # The template contains "XMTL-<###>" literally, so we replace the
    # placeholder with the bare numeric portion only.
    xmtl_bare = _normalize_xmtl_num(fields.get("transmittal_num", ""))

    # Text replacements
    mapping = {
        "<DATE>": fields.get("date", ""),
        "R3P-<PRJ#>": fields.get("job_num", ""),
        "XMTL-<###>": f"XMTL-{xmtl_bare}" if xmtl_bare else "",
        "<CLIENT> - <SITE NAME>": fields.get("client", ""),
        "<PROJECT DESCRIPTION>": fields.get("project_desc", ""),
        "Andrew Simmons, P.E.": fields.get("from_name", ""),
        "Managing Partner": fields.get("from_title", ""),
        "e: andrew.simmons@root3power.com": f"e: {fields.get('from_email', '')}",
        "c: (713) 294-2003": f"c: {fields.get('from_phone', '')}",
        "TX FIRM #20290": fields.get("firm", ""),
    }
    replace_text_everywhere(doc, mapping)

    # Checkboxes
    for label, key in CHECKBOX_MAP:
        set_checkbox(doc, label, checks.get(key, False))

    # Contacts
    fill_contacts_table(doc, contacts)

    # If no source PDFs are being included, strip the Reference heading +
    # its table BEFORE locating the Document Index table — both share the
    # same header row ("Document No." / "Description" / "Revision") and
    # find_table_by_headers returns the first match in document order.
    if not has_attached_drawings:
        _remove_reference_section(doc)

    # Document index
    idx_table = find_table_by_headers(doc, ["Document No.", "Description", "Revision"])
    if idx_table:
        clear_table_body(idx_table)
        append_rows(idx_table, documents)

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    return out_path
